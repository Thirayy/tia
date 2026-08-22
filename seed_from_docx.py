import os
import re
import pandas as pd
from docx import Document
from sqlmodel import Session, select, SQLModel
from app.database import engine
from app.models import TajwidRubrik  # Pakai model dari app.models

file_path = "Rubrik_Penilaian_Tajwid.docx"

def parse_docx_to_dict_list(path: str):
    """Fungsi helper untuk mengekstrak tabel .docx ke bentuk list dictionary"""
    if not os.path.exists(path):
        print(f"❌ File '{path}' tidak ditemukan!")
        return []
    
    doc = Document(path)
    parsed_data = []
    current_kategori = "Umum"

    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells):
                    continue
                
                # Deteksi baris Kategori Utama (misal: "A. Hukum Mad")
                if re.match(r"^[A-Z]\.\s+.+", cells[0]) or "Kaidah / Sub-Kaidah" in cells[0]:
                    if "Kaidah / Sub-Kaidah" not in cells[0]:
                        current_kategori = cells[0]
                    continue
                
                # Mapping kolom tabel ke key dictionary
                if len(cells) >= 3 and cells[0] and cells[0] != "Kaidah / Sub-Kaidah":
                    parsed_data.append({
                        "kategori": current_kategori,
                        "sub_kaidah": cells[0],
                        "keterangan": cells[1],
                        "kriteria_penilaian": cells[2]
                    })
    return parsed_data

def buat_tabel_dan_seed():
    print("🔄 Membuat tabel tajwid_rubrik di database (jika belum ada)...")
    SQLModel.metadata.create_all(engine)
    
    print(f"🔄 Lagi membaca file lokal {file_path}...")
    try:
        raw_data = parse_docx_to_dict_list(file_path)
        if not raw_data:
            print("⚠️ Tidak ada data yang berhasil diekstrak.")
            return

        # Masukkan ke Pandas DataFrame
        df = pd.DataFrame(raw_data)
        
        # Cleaning data NaN/Kosong dari DataFrame
        df['kategori'] = df['kategori'].fillna('Umum')
        df['sub_kaidah'] = df['sub_kaidah'].fillna('')
        df['keterangan'] = df['keterangan'].fillna('Tidak ada keterangan.')
        df['kriteria_penilaian'] = df['kriteria_penilaian'].fillna('Tidak ada kriteria.')
        
        total_rows = len(df)
        print(f"✅ File berhasil dibaca! Menemukan {total_rows} kaidah tajwid.")
        print("🚀 Memulai proses seeding ke database Postgres, tunggu bentar...")
        
        with Session(engine) as session:
            # Cek dulu biar gak double seed jika pernah dijalankan
            cek_data = session.exec(select(TajwidRubrik)).first()
            if cek_data:
                print("⚠️ Database udah ada isinya! Seeding dibatalkan biar gak double.")
                return

            for index, row in df.iterrows():
                # Mapping data dari DataFrame ke Model TajwidRubrik
                data_item = TajwidRubrik(
                    kategori=str(row['kategori']).strip(),
                    sub_kaidah=str(row['sub_kaidah']).strip(),
                    keterangan=str(row['keterangan']).strip(),
                    kriteria_penilaian=str(row['kriteria_penilaian']).strip()
                )
                session.add(data_item)
                
                # Commit bertahap per 20 baris
                if index % 20 == 0 and index > 0:
                    session.commit()
                    print(f"▓ {index}/{total_rows} kaidah tajwid berhasil masuk database...")
            
            # Commit sisa data
            session.commit()
            print("\n🎉 KELAR! Semua kaidah tajwid dari file .docx udah masuk ke Postgres!")

    except Exception as e:
        print("Waduh gagal, error-nya:", e)

if __name__ == "__main__":
    buat_tabel_dan_seed()